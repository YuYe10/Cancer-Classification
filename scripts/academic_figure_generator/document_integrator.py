"""
Document Integrator Module
======================

Provides functionality for integrating generated figures into academic
documents, supporting both Microsoft Word (.docx) and LaTeX (.tex) formats.

Author: Cancer-Classification Project
"""

from pathlib import Path
from typing import Optional, Union, List, Dict, Any, Tuple
from dataclasses import dataclass
import shutil
import re
import datetime


@dataclass
class FigureReference:
    """
    Represents a figure reference for document integration.

    Attributes:
        figure_id: Unique identifier for the figure.
        label: LaTeX label string.
        caption: Figure caption text.
        path: Path to the figure file.
        width: Figure width in the document (for LaTeX, as fraction of text width).
        position: Preferred position in document ('h', 't', 'b', 'p' for LaTeX).
    """
    figure_id: str
    label: str
    caption: str
    path: Path
    width: float = 0.9
    position: str = 'htbp'


@dataclass
class TableReference:
    """
    Represents a table reference for document integration.

    Attributes:
        table_id: Unique identifier for the table.
        label: LaTeX label string.
        caption: Table caption text.
        position: Preferred position in document.
    """
    table_id: str
    label: str
    caption: str
    position: str = 'htbp'


class LaTeXIntegrator:
    """
    Handles integration of figures into LaTeX documents.

    This class provides methods for inserting figures into LaTeX documents
    with proper formatting, captions, labels, and cross-references.

    Example:
        >>> integrator = LaTeXIntegrator('paper.tex')
        >>> integrator.add_figure(
        ...     figure_id='fig:performance',
        ...     caption='Model performance comparison',
        ...     path='outputs/figures/performance.pdf',
        ...     label='fig:performance'
        ... )
        >>> integrator.update_document()
    """

    LATEX_TEMPLATE = r"""
\\begin{{figure}}[{position}]
    \\centering
    \\includegraphics[width={width}\\textwidth]{{{path}}}
    \\caption{{{caption}}}
    \\label{{{label}}}
\\end{{figure}}
"""

    LATEX_SUBFIGURE_TEMPLATE = r"""
\\begin{{figure}}[{position}]
    \\centering
    \\begin{{subfigure}}[b]{{0.48\\textwidth}}
        \\centering
        \\includegraphics[width=\\textwidth]{{{path1}}}
        \\caption{{{caption1}}}
        \\label{{{label1}}}
    \\end{{subfigure}}
    \\hfill
    \\begin{{subfigure}}[b]{{0.48\\textwidth}}
        \\centering
        \\includegraphics[width=\\textwidth]{{{path2}}}
        \\caption{{{caption2}}}
        \\label{{{label2}}}
    \\end{{subfigure}}
    \\caption{{{main_caption}}}
    \\label{{{main_label}}}
\\end{{figure}}
"""

    def __init__(self, document_path: str):
        """
        Initialize the LaTeX integrator.

        Args:
            document_path: Path to the target LaTeX document.
        """
        self.document_path = Path(document_path)
        self.figures: List[FigureReference] = []
        self._original_content: Optional[str] = None

        if self.document_path.exists():
            with open(self.document_path, 'r', encoding='utf-8') as f:
                self._original_content = f.read()

    def add_figure(self,
                  figure_id: str,
                  caption: str,
                  path: Union[str, Path],
                  label: Optional[str] = None,
                  width: float = 0.9,
                  position: str = 'htbp',
                  insert_after: Optional[str] = None,
                  insert_before: Optional[str] = None) -> 'LaTeXIntegrator':
        """
        Add a figure to the document.

        Args:
            figure_id: Unique identifier for the figure.
            caption: Figure caption text.
            path: Path to the figure file.
            label: LaTeX label (auto-generated if None).
            width: Figure width as fraction of text width.
            position: LaTeX float position specifier.
            insert_after: Section or label to insert after.
            insert_before: Section or label to insert before.

        Returns:
            Self for method chaining.
        """
        if label is None:
            label = f"fig:{figure_id}"

        path = Path(path)
        if not path.is_absolute():
            path = path.resolve()

        figure_ref = FigureReference(
            figure_id=figure_id,
            label=label,
            caption=caption,
            path=path,
            width=width,
            position=position
        )

        self.figures.append(figure_ref)
        return self

    def add_figure_pair(self,
                       figure_id: str,
                       caption1: str,
                       path1: Union[str, Path],
                       caption2: str,
                       path2: Union[str, Path],
                       main_caption: str,
                       label: Optional[str] = None,
                       width: float = 0.48,
                       position: str = 'htbp') -> 'LaTeXIntegrator':
        """
        Add a pair of subfigures side by side.

        Args:
            figure_id: Unique identifier for the figure pair.
            caption1: Caption for the first subfigure.
            path1: Path to the first figure file.
            caption2: Caption for the second subfigure.
            path2: Path to the second figure file.
            main_caption: Main caption for the figure pair.
            label: LaTeX label (auto-generated if None).
            width: Individual subfigure width.
            position: LaTeX float position specifier.

        Returns:
            Self for method chaining.
        """
        if label is None:
            label = f"fig:{figure_id}"

        path1 = Path(path1) if not Path(path1).is_absolute() else path1
        path2 = Path(path2) if not Path(path2).is_absolute() else path2

        figure_content = self.LATEX_SUBFIGURE_TEMPLATE.format(
            position=position,
            path1=str(path1),
            caption1=caption1,
            label1=f"{label}a",
            path2=str(path2),
            caption2=caption2,
            label2=f"{label}b",
            main_caption=main_caption,
            main_label=label
        )

        self.figures.append(FigureReference(
            figure_id=figure_id,
            label=label,
            caption=main_caption,
            path=Path(""),
            width=width,
            position=position
        ))

        self._insert_content(figure_content)
        return self

    def _insert_content(self, content: str,
                       insert_after: Optional[str] = None,
                       insert_before: Optional[str] = None) -> None:
        """
        Insert content into the document at the specified location.

        Args:
            content: LaTeX content to insert.
            insert_after: Pattern to insert after.
            insert_before: Pattern to insert before.
        """
        if self._original_content is None:
            return

        if insert_after:
            pattern = re.compile(f'({re.escape(insert_after)})')
            match = pattern.search(self._original_content)
            if match:
                insert_pos = match.end()
                self._original_content = (
                    self._original_content[:insert_pos] +
                    '\n\n' + content + '\n\n' +
                    self._original_content[insert_pos:]
                )
        elif insert_before:
            pattern = re.compile(f'({re.escape(insert_before)})')
            match = pattern.search(self._original_content)
            if match:
                insert_pos = match.start()
                self._original_content = (
                    self._original_content[:insert_pos] +
                    '\n\n' + content + '\n\n' +
                    self._original_content[insert_pos:]
                )
        else:
            self._original_content += '\n\n' + content + '\n\n'

    def update_document(self,
                      backup: bool = True,
                      insert_mode: str = 'end') -> Path:
        """
        Update the LaTeX document with all added figures.

        Args:
            backup: Whether to create a backup of the original.
            insert_mode: Where to insert figures ('end', 'section', 'begin').

        Returns:
            Path to the updated document.
        """
        if self._original_content is None:
            raise ValueError("No document loaded")

        if backup:
            backup_path = self.document_path.with_suffix('.tex.bak')
            shutil.copy(self.document_path, backup_path)

        if insert_mode == 'end':
            pass
        elif insert_mode == 'begin':
            for fig in reversed(self.figures):
                if fig.path != Path(""):
                    figure_content = self.LATEX_TEMPLATE.format(
                        position=fig.position,
                        width=fig.width,
                        path=str(fig.path),
                        caption=fig.caption,
                        label=fig.label
                    )
                    self._original_content = figure_content + '\n\n' + self._original_content

        with open(self.document_path, 'w', encoding='utf-8') as f:
            f.write(self._original_content)

        return self.document_path

    def generate_figure_listing(self) -> str:
        """
        Generate a list of all figures with their references.

        Returns:
            LaTeX code for a figure listing.
        """
        listing = "\\listoffigures\n\\newpage\n\n"

        for fig in self.figures:
            listing += f"\\ref{{{fig.label}}} - {fig.caption}\n"

        return listing

    def generate_tex_snippet(self, figure: FigureReference) -> str:
        """
        Generate LaTeX code for a single figure.

        Args:
            figure: Figure reference to generate code for.

        Returns:
            LaTeX code string.
        """
        if figure.path == Path(""):
            return ""

        return self.LATEX_TEMPLATE.format(
            position=figure.position,
            width=figure.width,
            path=str(figure.path),
            caption=figure.caption,
            label=figure.label
        )


class WordIntegrator:
    """
    Handles integration of figures into Microsoft Word documents.

    This class provides methods for inserting figures into Word documents
    with proper formatting, captions, and cross-references.

    Note:
        Requires python-docx package for Word document handling.

    Example:
        >>> integrator = WordIntegrator('paper.docx')
        >>> integrator.add_figure(
        ...     caption='Model performance comparison',
        ...     path='outputs/figures/performance.png',
        ...     width_inches=6.0
        ... )
        >>> integrator.update_document()
    """

    def __init__(self, document_path: str):
        """
        Initialize the Word integrator.

        Args:
            document_path: Path to the target Word document.
        """
        self.document_path = Path(document_path)
        self.figures: List[Dict[str, Any]] = []

        try:
            from docx import Document
            self.doc = Document(document_path)
        except ImportError:
            raise ImportError("python-docx is required for Word integration. "
                            "Install with: pip install python-docx")

    def add_figure(self,
                   caption: str,
                   path: Union[str, Path],
                   width_inches: float = 6.0,
                   height_inches: Optional[float] = None,
                   center: bool = True) -> 'WordIntegrator':
        """
        Add a figure to the document.

        Args:
            caption: Figure caption text.
            path: Path to the figure file.
            width_inches: Figure width in inches.
            height_inches: Figure height in inches (auto-calculated if None).
            center: Whether to center the figure.

        Returns:
            Self for method chaining.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Figure not found: {path}")

        self.figures.append({
            'caption': caption,
            'path': path,
            'width': width_inches,
            'height': height_inches,
            'center': center
        })

        return self

    def update_document(self,
                      output_path: Optional[str] = None) -> Path:
        """
        Update the Word document with all added figures.

        Args:
            output_path: Path for the output document.
                        Uses original path if None.

        Returns:
            Path to the updated document.
        """
        for fig in self.figures:
            if fig['center']:
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = 1

            run = self.doc.add_picture(str(fig['path']), width=fig['width'])

            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.add_run(f"\n\nFigure: {fig['caption']}")

        output_path = output_path or str(self.document_path)

        self.doc.save(output_path)

        return Path(output_path)


class DocumentIntegrator:
    """
    Unified interface for document integration supporting both
    LaTeX and Word formats.

    This class provides a high-level interface for integrating
    figures into academic documents, automatically detecting
    the document format and using the appropriate integrator.

    Example:
        >>> integrator = DocumentIntegrator('paper.tex')
        >>> integrator.add_figure(
        ...     caption='Experimental results',
        ...     path='outputs/figures/results.pdf'
        ... )
        >>> integrator.update()
    """

    def __init__(self, document_path: str):
        """
        Initialize the document integrator.

        Args:
            document_path: Path to the target document.
        """
        self.document_path = Path(document_path)
        suffix = self.document_path.suffix.lower()

        if suffix == '.tex':
            self._integrator = LaTeXIntegrator(document_path)
        elif suffix == '.docx':
            self._integrator = WordIntegrator(document_path)
        else:
            raise ValueError(f"Unsupported document format: {suffix}")

    def add_figure(self,
                  figure_id: str,
                  caption: str,
                  path: Union[str, Path],
                  label: Optional[str] = None,
                  width: float = 0.9,
                  position: str = 'htbp') -> 'DocumentIntegrator':
        """
        Add a figure to the document.

        Args:
            figure_id: Unique identifier for the figure.
            caption: Figure caption text.
            path: Path to the figure file.
            label: Label for cross-referencing.
            width: Figure width (LaTeX) or height in inches (Word).
            position: Preferred position in document.

        Returns:
            Self for method chaining.
        """
        if isinstance(self._integrator, LaTeXIntegrator):
            self._integrator.add_figure(
                figure_id=figure_id,
                caption=caption,
                path=path,
                label=label,
                width=width,
                position=position
            )
        elif isinstance(self._integrator, WordIntegrator):
            self._integrator.add_figure(
                caption=caption,
                path=path,
                width_inches=width * 10
            )

        return self

    def update(self, backup: bool = True) -> Path:
        """
        Update the document with all added figures.

        Args:
            backup: Whether to create a backup (LaTeX only).

        Returns:
            Path to the updated document.
        """
        if isinstance(self._integrator, LaTeXIntegrator):
            return self._integrator.update_document(backup=backup)
        else:
            return self._integrator.update_document()

    def generate_figure_list(self) -> str:
        """
        Generate a list of all figures with references.

        Returns:
            LaTeX code for figure listing if LaTeX document.
        """
        if isinstance(self._integrator, LaTeXIntegrator):
            return self._integrator.generate_figure_listing()
        return ""

    @property
    def figures(self) -> List:
        """Return the list of added figures."""
        return self._integrator.figures

    def __repr__(self) -> str:
        return f"DocumentIntegrator(path='{self.document_path}')"